#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
智刷星 · 本地服务 + LLM 代理
================================
- 静态文件服务（取代原来的 python -m http.server）
- POST /api/analyze : 调用 OpenAI 兼容的大模型，对用户粘贴的题目做
  知识点归纳与复习总结。
- POST /api/parse-file : 上传 PDF / Word(.docx) / TXT 文件，服务端抽取纯文本
  （PDF 依赖 pypdf，Word 依赖 python-docx；扫描件 PDF 自动改用本地离线 OCR）。
- POST /api/ocr : 上传题目图片（dataURL），优先使用本地离线 OCR 引擎
  （rapidocr-onnxruntime，无需任何 API Key）；若引擎不可用则回退到支持视觉的
  大模型（需配置 LLM_MODEL）。

安全说明：
- API Key 仅保存在本机（站点目录下的 .env），服务端读取，绝不下发到浏览器。
- 浏览器只把"题目文本" POST 到同源的 /api/analyze，拿回分析结果。
- 支持的厂商：任何 OpenAI 兼容接口，例如 OpenAI / DeepSeek / Moonshot(Kimi)
  / 智谱 GLM / 通义千问(Qwen) 等，只需在 .env 里改 LLM_BASE_URL 与 LLM_MODEL。
"""

import http.server
import json
import os
import re
import socketserver
import sys
import base64
import io
import threading
import time
import collections
import urllib.error
import urllib.request
from pathlib import Path

# 可选依赖（用于文件解析 / 视觉 OCR），缺失时相关接口会优雅降级而非崩溃
try:
    import pypdf
    HAVE_PYPDF = True
except Exception:  # noqa: BLE001
    HAVE_PYPDF = False
try:
    import docx
    HAVE_DOCX = True
except Exception:  # noqa: BLE001
    HAVE_DOCX = False

BASE_DIR = Path(__file__).resolve().parent
PORT = int(os.environ.get("PORT", "8123"))


# ----------------------------------------------------------------------
# 读取 .env（简单解析，不依赖第三方库）
# ----------------------------------------------------------------------
def load_env():
    env_path = BASE_DIR / ".env"
    if not env_path.exists():
        return
    for line in env_path.read_text(encoding="utf-8").splitlines():
        line = line.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        k, v = line.split("=", 1)
        k, v = k.strip(), v.strip()
        if len(v) >= 2 and v[0] in "\"'" and v[-1] == v[0]:
            v = v[1:-1]
        os.environ.setdefault(k, v)


load_env()

LLM_API_KEY = os.environ.get("LLM_API_KEY", "").strip()
LLM_BASE_URL = os.environ.get("LLM_BASE_URL", "https://api.openai.com/v1").strip().rstrip("/")
LLM_MODEL = os.environ.get("LLM_MODEL", "gpt-4o-mini").strip()
LLM_TIMEOUT = int(os.environ.get("LLM_TIMEOUT", "90"))
LLM_ENABLED = bool(LLM_API_KEY)
# 可选：公网部署时设置 SITE_TOKEN，所有 /api/* 接口必须带此口令（X-Site-Token 头或
# ?site_token= 参数），否则 403。防止公网链接泄露后被他人白嫖你的模型额度。
SITE_TOKEN = os.environ.get("SITE_TOKEN", "").strip()

# ----------------------------------------------------------------------
# 轻量级限流（保护公网部署时的模型额度，按客户端 IP + 接口，每分钟上限）
# ----------------------------------------------------------------------
_RATE_LIMIT = {"gen-exam": 15, "analyze": 30, "ocr": 20, "parse-file": 60}
_hits = collections.defaultdict(list)
_rl_lock = threading.Lock()


def _client_ip(handler):
    """取真实客户端 IP；穿透 Cloudflare / 反向代理时优先用转发头。"""
    fwd = handler.headers.get("X-Forwarded-For") or handler.headers.get("CF-Connecting-IP")
    if fwd:
        return fwd.split(",")[0].strip()
    return handler.client_address[0]


def check_rate(handler, route):
    ip = _client_ip(handler)
    now = time.time()
    key = (ip, route)
    with _rl_lock:
        lst = [t for t in _hits[key] if now - t < 60]
        _hits[key] = lst
        limit = _RATE_LIMIT.get(route, 30)
        if len(lst) >= limit:
            return False
        _hits[key].append(now)
    return True


def check_token(handler):
    if not SITE_TOKEN:
        return True
    provided = handler.headers.get("X-Site-Token") or ""
    if not provided and "?" in handler.path:
        q = handler.path.split("?", 1)[1]
        for pair in q.split("&"):
            if pair.startswith("site_token="):
                provided = pair.split("=", 1)[1]
                break
    return provided == SITE_TOKEN

SYSTEM_PROMPT = """你是一名资深大学课程助教，擅长从题目中抽取考察的知识点并给出复习建议。
用户会给你一批题目（可能来自不同科目、不同题型）。请完成两件事：

1. 逐题分析：对每道题，给出 1-4 个简洁、具体、可操作的知识点标签（中文短语，
   例如「条件概率与贝叶斯」「导数计算」「辛亥革命」「虚拟语气」），避免使用过于宽泛的词。
2. 整体总结：综合所有题目，给出
   - topics：检测到的主要专题（3-6 个短语）；
   - narrative：2-4 句自然语言总结，说明这批题覆盖的范围、重点，以及建议优先复习的方向；
   - reviewPlan：2-4 条具体复习建议（每条一句话）。

必须严格返回如下 JSON（不要包含任何额外文字，也不要使用 markdown 代码块）：
{
  "items": [ { "index": 0, "kps": ["知识点1", "知识点2"] } ],
  "summary": {
    "topics": ["专题A", "专题B"],
    "narrative": "……",
    "reviewPlan": ["建议1", "建议2"]
  }
}
其中 items 的顺序 index 与输入题目一致（从 0 开始）。若某题无法判断知识点，kps 给空数组。"""


def extract_json(text):
    """尽力从模型返回里解析出 JSON 对象。"""
    if text is None:
        return {}
    text = text.strip()
    # 去掉可能的 ```json ... ``` 代码围栏
    m = re.search(r"```(?:json)?\s*([\s\S]*?)```", text, re.IGNORECASE)
    if m:
        text = m.group(1).strip()
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass
    # 取第一个 { 到最后一个 } 之间的内容
    s, e = text.find("{"), text.rfind("}")
    if s != -1 and e != -1 and e > s:
        try:
            return json.loads(text[s:e + 1])
        except json.JSONDecodeError:
            return {}
    return {}


def call_llm(questions):
    url = f"{LLM_BASE_URL}/chat/completions"
    payload = {
        "model": LLM_MODEL,
        "messages": [
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": json.dumps({"questions": questions}, ensure_ascii=False)},
        ],
        "temperature": 0.2,
    }
    # 部分厂商支持 json 输出格式，加上更稳；不支持的会忽略
    if "openai.com" in LLM_BASE_URL or "deepseek" in LLM_BASE_URL or "moonshot" in LLM_BASE_URL:
        payload["response_format"] = {"type": "json_object"}

    data = json.dumps(payload).encode("utf-8")
    req = urllib.request.Request(
        url,
        data=data,
        headers={
            "Content-Type": "application/json",
            "Authorization": f"Bearer {LLM_API_KEY}",
        },
        method="POST",
    )
    with urllib.request.urlopen(req, timeout=LLM_TIMEOUT) as resp:
        body = json.loads(resp.read().decode("utf-8"))
    content = body["choices"][0]["message"]["content"]
    return extract_json(content)


# ----------------------------------------------------------------------
# 视觉 OCR：调用支持视觉的 OpenAI 兼容模型，从题目图片抽取纯文本
# ----------------------------------------------------------------------
VISION_PROMPT = """你是一名试卷 OCR 与题目抽取助手。用户会发来一张题目图片（可能是中文试卷，含数学公式）。
请忠实提取图中所有题目，并按如下纯文本格式输出，以便下游程序自动解析：

- 每题以“数字+标点”开头，例如 "1. " 或 "1、"。
- 选择题的四个选项写在同一行，用空格分隔：A. xxx  B. yyy  C. zzz  D. www。
- 填空题的空用 "____" 表示。
- 若图中出现“参考答案 / 答案”，请在末尾另起一行写：参考答案，然后按题号列出，例如：
参考答案
1. A 2. B 3. C
（若没有答案页，则不要编造，直接省略该部分）
- 数学符号尽量用纯文本保留，如 X~N(0,1)、P(A)=0.5、(X,Y)。
- 只输出提取出的题目文本，不要加任何解释，也不要使用 markdown 代码块。
- 如果图片不是题目/试卷，用一句话说明你看到的内容。"""


def call_vision_llm(data_url):
    """调用支持视觉的 OpenAI 兼容模型，从图片抽取题目文本。"""
    url = f"{LLM_BASE_URL}/chat/completions"
    payload = {
        "model": LLM_MODEL,
        "messages": [
            {"role": "system", "content": VISION_PROMPT},
            {"role": "user", "content": [
                {"type": "text", "text": "请提取这张图片里的全部题目。"},
                {"type": "image_url", "image_url": {"url": data_url}},
            ]},
        ],
        "temperature": 0.1,
    }
    data = json.dumps(payload).encode("utf-8")
    req = urllib.request.Request(
        url,
        data=data,
        headers={
            "Content-Type": "application/json",
            "Authorization": f"Bearer {LLM_API_KEY}",
        },
        method="POST",
    )
    with urllib.request.urlopen(req, timeout=LLM_TIMEOUT) as resp:
        body = json.loads(resp.read().decode("utf-8"))
    return body["choices"][0]["message"]["content"]


# ----------------------------------------------------------------------
# AI 生成模拟卷：根据识别出的知识点，调用大模型原创编写全新题目
# ----------------------------------------------------------------------
GEN_EXAM_PROMPT = """你是一名严谨的大学课程出题老师。请根据下面给出的「知识点清单」，原创编写一批全新的练习题。
重要：必须是原创新题，不得照抄或复述用户提供的原题（sample 仅作难度与风格参考）。

要求：
1. 题型混合，至少包含：选择题（4 个选项，answer 用正确选项字母）、填空题、以及计算题/证明题/应用题中的一种或多种。
2. 难度相当于大学对应课程的期中/期末水平，表述严谨；数学符号用纯文本，如 X~N(0,1)、P(A)=0.5、(X,Y)。
3. 题目数量严格等于用户要求的 count。
4. 每道题结构：
   - type: "选择题" | "填空题" | "计算题" | "证明题" | "应用题"
   - stem: 题干（填空题用 ____ 表示空）
   - options: 选择题为 ["A. ...","B. ...","C. ...","D. ..."]，其它题型为 []
   - answer: 选择题填选项字母（如 "B"），填空/大题填最终结论或数值
   - solution: 1-3 句关键解题步骤
   - kp: 该题对应的知识点（取自给定清单）
5. 只返回如下 JSON 对象，不要任何额外文字、不要 markdown 代码块：
{"questions":[ {"type":"选择题","stem":"...","options":["A. ...","B. ...","C. ...","D. ..."],"answer":"B","solution":"...","kp":"..."}, {"type":"填空题","stem":"... ____ ...","options":[],"answer":"...","solution":"...","kp":"..."}, {"type":"计算题","stem":"...","options":[],"answer":"...","solution":"...","kp":"..."} ]}"""


def gen_exam_llm(topics, count, course, sample):
    """调用大模型，按知识点原创生成 count 道新题；返回题目 dict 列表。"""
    url = f"{LLM_BASE_URL}/chat/completions"
    user_content = {
        "course": course or "（未知，请根据知识点自行判断所属学科与课程）",
        "topics": topics,
        "count": count,
        "sample": (sample or "")[:2000],
        "note": "围绕上述知识点原创出题，题型按上述要求混合；sample 仅作风格/难度参考，严禁抄袭原题。",
    }
    payload = {
        "model": LLM_MODEL,
        "messages": [
            {"role": "system", "content": GEN_EXAM_PROMPT},
            {"role": "user", "content": json.dumps(user_content, ensure_ascii=False)},
        ],
        "temperature": 0.7,
    }
    if "openai.com" in LLM_BASE_URL or "deepseek" in LLM_BASE_URL or "moonshot" in LLM_BASE_URL:
        payload["response_format"] = {"type": "json_object"}
    data = json.dumps(payload).encode("utf-8")
    req = urllib.request.Request(
        url,
        data=data,
        headers={
            "Content-Type": "application/json",
            "Authorization": f"Bearer {LLM_API_KEY}",
        },
        method="POST",
    )
    with urllib.request.urlopen(req, timeout=LLM_TIMEOUT) as resp:
        body = json.loads(resp.read().decode("utf-8"))
    content = body["choices"][0]["message"]["content"]
    parsed = extract_json(content)
    if isinstance(parsed, dict):
        for k in ("questions", "items", "data", "exams"):
            if isinstance(parsed.get(k), list):
                return parsed[k]
    if isinstance(parsed, list):
        return parsed
    return []


# ----------------------------------------------------------------------
# 本地离线 OCR：使用 rapidocr-onnxruntime（纯 Python，无需任何 API Key）
# 用于扫描件 PDF 与图片上传；优先于需要联网密钥的视觉大模型。
# ----------------------------------------------------------------------
_OCR_ENGINE = None
_ocr_lock = threading.Lock()


def _ocr_importable():
    try:
        import rapidocr_onnxruntime  # noqa: F401
        return True
    except Exception:  # noqa: BLE001
        return False


OCR_AVAILABLE = _ocr_importable()


def _get_ocr():
    """懒加载并缓存 RapidOCR 实例；不可用时返回 None。"""
    global _OCR_ENGINE
    if _OCR_ENGINE is not None:
        return _OCR_ENGINE if _OCR_ENGINE is not False else None
    with _ocr_lock:
        if _OCR_ENGINE is None:
            try:
                from rapidocr_onnxruntime import RapidOCR
                _OCR_ENGINE = RapidOCR()
            except Exception:  # noqa: BLE001
                _OCR_ENGINE = False  # 标记为不可用，避免反复尝试
    return _OCR_ENGINE if _OCR_ENGINE is not False else None


def ocr_image_bytes(img_bytes):
    """对单张图片做离线 OCR，返回按行拼接的文本；失败返回空串。"""
    engine = _get_ocr()
    if engine is None:
        return ""
    try:
        import numpy as np
        from PIL import Image
        img = Image.open(io.BytesIO(img_bytes)).convert("RGB")
        arr = np.asarray(img)
        with _ocr_lock:
            result, _ = engine(arr)
        if not result:
            return ""
        return "\n".join(box[1] for box in result)
    except Exception:  # noqa: BLE001
        return ""


def _largest_page_image(pdf_page):
    """从 PDF 页面取面积最大的内嵌图片字节（即整页扫描图）。"""
    try:
        images = pdf_page.images
    except Exception:  # noqa: BLE001
        return None
    best, best_area = None, 0
    for im in images:
        try:
            from PIL import Image as _PIL
            area = _PIL.open(io.BytesIO(im.data)).size
            area = area[0] * area[1]
        except Exception:  # noqa: BLE001
            area = len(im.data)
        if area > best_area:
            best_area, best = area, im.data
    return best


def ocr_pdf(pdf_bytes):
    """对扫描件 PDF：逐页抽取整页图并本地 OCR；返回 (text, note)。"""
    if not HAVE_PYPDF:
        return "", "服务端未安装 pypdf。"
    try:
        reader = pypdf.PdfReader(io.BytesIO(pdf_bytes))
    except Exception as e:  # noqa: BLE001
        return "", f"PDF 解析失败：{e}"
    if not OCR_AVAILABLE:
        return "", ("未安装本地 OCR 引擎（rapidocr-onnxruntime），无法识别扫描件 PDF。"
                    "请运行 pip install rapidocr-onnxruntime，或改用「上传图片」+ 视觉大模型。")
    pages_text = []
    for i, pg in enumerate(reader.pages):
        img = _largest_page_image(pg)
        if not img:
            continue
        t = ocr_image_bytes(img)
        if t.strip():
            pages_text.append(f"===== 第{i + 1}页 =====\n{t}")
    text = "\n\n".join(pages_text).strip()
    note = "扫描件 PDF：已使用本地离线 OCR 逐页识别（无需联网密钥）。" if text else \
        "本地 OCR 未能从页面图片中提取到文字。"
    return text, note


def extract_file_text(filename, data):
    """根据扩展名抽取文本；扫描件 PDF 自动改用本地离线 OCR。"""
    note = ""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext == "pdf":
        if not HAVE_PYPDF:
            return "", "服务端未安装 pypdf，无法解析 PDF（请运行 pip install pypdf）。"
        try:
            reader = pypdf.PdfReader(io.BytesIO(data))
            parts = [ (pg.extract_text() or "") for pg in reader.pages ]
            text = "\n".join(parts).strip()
        except Exception as e:  # noqa: BLE001
            return "", f"PDF 解析失败：{e}"
        # 清洗明显的扫描水印（如「扫描全能王 创建」）后判断文字层是否过薄
        cleaned = "\n".join(l for l in text.splitlines()
                            if "扫描全能王" not in l and l.strip())
        if len(cleaned) < 120:
            if OCR_AVAILABLE:
                ocr_text, ocr_note = ocr_pdf(data)
                if ocr_text:
                    return ocr_text, ocr_note
                return text, (ocr_note or "扫描件 PDF 已尝试本地 OCR，但未提取到文字。")
            note = ("该 PDF 为扫描图片型，没有可提取的文字层，且服务端未安装本地 OCR 引擎。"
                    "请运行 pip install rapidocr-onnxruntime，或改用「上传图片」+ 视觉大模型识别。")
            return text, note
        return text, note
    if ext == "docx":
        if not HAVE_DOCX:
            return "", "服务端未安装 python-docx，无法解析 Word（请运行 pip install python-docx）。"
        try:
            doc = docx.Document(io.BytesIO(data))
            paras = [p.text for p in doc.paragraphs]
            for tbl in doc.tables:
                for row in tbl.rows:
                    paras.append(" | ".join(c.text for c in row.cells))
            text = "\n".join([p for p in paras if p.strip()]).strip()
        except Exception as e:  # noqa: BLE001
            return "", f"Word 解析失败：{e}"
        return text, note
    if ext in ("txt", "md", "csv", "text"):
        for enc in ("utf-8", "gb18030", "latin-1"):
            try:
                return data.decode(enc).strip(), note
            except Exception:  # noqa: BLE001
                continue
        return "", "文本解码失败（尝试了 utf-8 / gb18030 / latin-1）。"
    return "", f"暂不支持的文件类型：.{ext}（支持 pdf / docx / txt / md / csv）。"


class Handler(http.server.BaseHTTPRequestHandler):
    def _send_json(self, obj, code=200):
        body = json.dumps(obj, ensure_ascii=False).encode("utf-8")
        self.send_response(code)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def do_POST(self):
        route = self.path.split("?")[0].lstrip("/")
        if route.startswith("api/"):
            if not check_token(self):
                return self._send_json({"error": "token required", "tokenRequired": True}, 403)
            rname = route.split("/")[-1]
            if not check_rate(self, rname):
                return self._send_json(
                    {"error": "请求过于频繁，请稍后再试（已启用限流以保护你的模型额度）。"}, 429)
        if self.path.split("?")[0] == "/api/analyze":
            try:
                length = int(self.headers.get("Content-Length", 0) or 0)
                raw = self.rfile.read(length) if length else b"{}"
                incoming = json.loads(raw.decode("utf-8") or "{}")
            except Exception as e:  # noqa: BLE001
                return self._send_json({"llmEnabled": False, "error": f"请求解析失败: {e}"}, 400)

            if not LLM_ENABLED:
                return self._send_json({
                    "llmEnabled": False,
                    "reason": "LLM 未配置：站点目录下缺少 .env 或未设置 LLM_API_KEY。",
                })

            questions = incoming.get("questions", [])
            if not questions:
                return self._send_json({
                    "llmEnabled": True,
                    "items": [],
                    "summary": {"topics": [], "narrative": "", "reviewPlan": []},
                })

            try:
                result = call_llm(questions)
                result["llmEnabled"] = True
                self._send_json(result)
            except urllib.error.HTTPError as e:
                detail = e.read().decode("utf-8", "ignore")[:300] if e.fp else str(e)
                self._send_json({"llmEnabled": False, "error": f"LLM 接口返回错误 {e.code}: {detail}"})
            except Exception as e:  # noqa: BLE001
                self._send_json({"llmEnabled": False, "error": f"LLM 调用失败: {e}"})
            return

        # ---- 文件解析：PDF / Word / TXT，返回抽取出的纯文本 ----
        if self.path.split("?")[0] == "/api/parse-file":
            try:
                length = int(self.headers.get("Content-Length", 0) or 0)
                raw = self.rfile.read(length) if length else b"{}"
                incoming = json.loads(raw.decode("utf-8") or "{}")
            except Exception as e:  # noqa: BLE001
                return self._send_json({"ok": False, "error": f"请求解析失败: {e}"}, 400)
            fname = (incoming.get("filename") or "file").lower()
            b64 = incoming.get("data") or ""
            if not b64:
                return self._send_json({"ok": False, "error": "缺少文件数据 data。"})
            if "," in b64 and b64[:20].lower().startswith("data:"):
                b64 = b64.split(",", 1)[1]  # 去掉 dataURL 头
            try:
                file_bytes = base64.b64decode(b64)
            except Exception as e:  # noqa: BLE001
                return self._send_json({"ok": False, "error": f"Base64 解码失败: {e}"})
            if len(file_bytes) > 15 * 1024 * 1024:
                return self._send_json({"ok": False, "error": "文件过大（上限 15MB）。"})
            text, note = extract_file_text(fname, file_bytes)
            return self._send_json({"ok": bool(text), "text": text, "note": note})

        # ---- 图片 OCR：优先本地离线 OCR（无需 Key），否则视觉大模型 ----
        if self.path.split("?")[0] == "/api/ocr":
            try:
                length = int(self.headers.get("Content-Length", 0) or 0)
                raw = self.rfile.read(length) if length else b"{}"
                incoming = json.loads(raw.decode("utf-8") or "{}")
            except Exception as e:  # noqa: BLE001
                return self._send_json({"ok": False, "error": f"请求解析失败: {e}"}, 400)
            image = incoming.get("image") or ""
            if not image:
                return self._send_json({"ok": False, "error": "缺少图片数据 image。"})
            # 兼容 dataURL 与裸 base64
            if "," in image and image[:20].lower().startswith("data:"):
                b64 = image.split(",", 1)[1]
            else:
                b64 = image
            try:
                img_bytes = base64.b64decode(b64)
            except Exception as e:  # noqa: BLE001
                return self._send_json({"ok": False, "error": f"图片解码失败: {e}"})
            # 1) 本地离线 OCR（无需任何 API Key）
            engine = _get_ocr()
            if engine is not None:
                text = ocr_image_bytes(img_bytes)
                if text.strip():
                    return self._send_json({"ok": True, "text": text, "ocrMode": "local"})
                return self._send_json({"ok": False, "ocrMode": "local",
                                        "error": "本地 OCR 未能从该图片提取到文字（可能是空白/模糊/非题目图片）。"})
            # 2) 回退：视觉大模型（需配置支持视觉的 LLM_MODEL）
            if not LLM_ENABLED:
                return self._send_json({
                    "ok": False,
                    "ocrEnabled": False,
                    "reason": "图片识别不可用：本地 OCR 引擎未安装，且未配置支持视觉的大模型。"
                              "请运行 pip install rapidocr-onnxruntime，或在 .env 设置 "
                              "LLM_API_KEY / LLM_BASE_URL / LLM_MODEL（支持视觉的模型）。",
                })
            try:
                text = call_vision_llm(image)
                self._send_json({"ok": True, "text": text or "", "ocrMode": "vision"})
            except urllib.error.HTTPError as e:
                detail = e.read().decode("utf-8", "ignore")[:400] if e.fp else str(e)
                self._send_json({"ok": False, "ocrMode": "vision",
                                 "error": f"视觉模型返回错误 {e.code}：{detail}（请确认所用模型支持视觉，如 qwen-vl / gpt-4o / moonshot-vision）。"})
            except Exception as e:  # noqa: BLE001
                self._send_json({"ok": False, "ocrMode": "vision", "error": f"视觉识别失败：{e}"})
            return

        # ---- AI 生成模拟卷：按知识点原创出题（需配置 LLM Key）----
        if self.path.split("?")[0] == "/api/gen-exam":
            try:
                length = int(self.headers.get("Content-Length", 0) or 0)
                raw = self.rfile.read(length) if length else b"{}"
                incoming = json.loads(raw.decode("utf-8") or "{}")
            except Exception as e:  # noqa: BLE001
                return self._send_json({"ok": False, "error": f"请求解析失败: {e}"}, 400)
            if not LLM_ENABLED:
                return self._send_json({
                    "ok": False,
                    "enabled": False,
                    "reason": "AI 生成未启用：站点目录下缺少 .env 或未设置 LLM_API_KEY。"
                              "请在 .env 配置 LLM_API_KEY / LLM_BASE_URL / LLM_MODEL 后重启服务。",
                })
            topics = incoming.get("topics") or []
            count = int(incoming.get("count", 6))
            course = (incoming.get("course") or "")[:40]
            sample = incoming.get("sample") or ""
            if not topics:
                return self._send_json({"ok": False, "error": "缺少知识点清单 topics（请先解析题目识别知识点）。"})
            try:
                questions = gen_exam_llm(topics, count, course, sample)
                if not questions:
                    return self._send_json({"ok": False, "error": "模型未返回有效题目，请重试或更换模型。"})
                self._send_json({"ok": True, "enabled": True, "questions": questions, "model": LLM_MODEL})
            except urllib.error.HTTPError as e:
                detail = e.read().decode("utf-8", "ignore")[:400] if e.fp else str(e)
                self._send_json({"ok": False, "error": f"LLM 接口返回错误 {e.code}: {detail}"})
            except Exception as e:  # noqa: BLE001
                self._send_json({"ok": False, "error": f"AI 生成失败: {e}"})
            return

        self._send_json({"error": "not found"}, 404)

    def do_GET(self):
        path = self.path.split("?")[0]
        if path == "/":
            path = "/index.html"
        target = (BASE_DIR / path.lstrip("/")).resolve()
        # 禁止目录穿越
        if not str(target).startswith(str(BASE_DIR)):
            self.send_error(403)
            return
        # 安全：禁止直接访问点文件（如 .env）与服务端脚本，避免泄露密钥/源码
        if target.name.startswith(".") or target.name in ("server.py",):
            self.send_error(403)
            return
        if target.is_dir():
            target = target / "index.html"
        if not target.exists():
            self.send_error(404)
            return
        # 仅允许静态资源类型，其余一律拒绝
        ALLOWED = {".html", ".css", ".js", ".json", ".svg", ".png",
                   ".jpg", ".jpeg", ".ico", ".woff", ".woff2", ".webp"}
        if target.suffix.lower() not in ALLOWED:
            self.send_error(403)
            return
        ctype = {
            ".html": "text/html; charset=utf-8",
            ".css": "text/css; charset=utf-8",
            ".js": "application/javascript; charset=utf-8",
            ".json": "application/json; charset=utf-8",
            ".svg": "image/svg+xml",
            ".png": "image/png",
            ".jpg": "image/jpeg",
            ".ico": "image/x-icon",
        }.get(target.suffix.lower(), "application/octet-stream")
        try:
            body = target.read_bytes()
            self.send_response(200)
            self.send_header("Content-Type", ctype)
            self.send_header("Content-Length", str(len(body)))
            self.end_headers()
            self.wfile.write(body)
        except Exception:  # noqa: BLE001
            self.send_error(500)

    def log_message(self, fmt, *args):
        sys.stderr.write("%s - %s\n" % (self.address_string(), fmt % args))


if __name__ == "__main__":
    socketserver.ThreadingTCPServer.allow_reuse_address = True
    # 监听 0.0.0.0：既可在本机访问，也可被隧道/云主机映射到公网。
    # 仅本机使用时同样正常（localhost 仍可达）。
    HOST = os.environ.get("BIND_HOST", "0.0.0.0")
    with socketserver.ThreadingTCPServer((HOST, PORT), Handler) as httpd:
        print(f"智刷星 服务已启动: http://{HOST}:{PORT}")
        deps = []
        deps.append("LLM=" + (LLM_MODEL if LLM_ENABLED else "未配置"))
        deps.append("pypdf=" + ("有" if HAVE_PYPDF else "无"))
        deps.append("docx=" + ("有" if HAVE_DOCX else "无"))
        deps.append("本地OCR=" + ("有(rapidocr)" if OCR_AVAILABLE else "无"))
        print("依赖状态: " + " | ".join(deps))
        try:
            httpd.serve_forever()
        except KeyboardInterrupt:
            print("\n已停止。")
